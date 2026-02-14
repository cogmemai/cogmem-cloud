"""Document parsing utilities for PDF and DOCX files."""

from __future__ import annotations

import io
import logging
from dataclasses import dataclass

logger = logging.getLogger(__name__)


@dataclass
class ParsedDocument:
    title: str
    text: str
    page_count: int
    word_count: int
    content_type: str


def parse_pdf(file_bytes: bytes, filename: str) -> ParsedDocument:
    """Extract text from a PDF file using pymupdf."""
    import fitz  # pymupdf

    doc = fitz.open(stream=file_bytes, filetype="pdf")
    pages: list[str] = []
    for page in doc:
        pages.append(page.get_text())
    doc.close()

    text = "\n\n".join(pages).strip()
    return ParsedDocument(
        title=filename,
        text=text,
        page_count=len(pages),
        word_count=len(text.split()),
        content_type="application/pdf",
    )


def parse_docx(file_bytes: bytes, filename: str) -> ParsedDocument:
    """Extract text from a DOCX file using python-docx."""
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text = "\n\n".join(paragraphs).strip()

    return ParsedDocument(
        title=filename,
        text=text,
        page_count=len(paragraphs),
        word_count=len(text.split()),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


def parse_document(file_bytes: bytes, filename: str, content_type: str) -> ParsedDocument:
    """Parse a document based on its content type or filename extension."""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if content_type == "application/pdf" or ext == "pdf":
        return parse_pdf(file_bytes, filename)
    elif ext in ("docx",) or "wordprocessingml" in content_type:
        return parse_docx(file_bytes, filename)
    elif ext == "txt" or content_type.startswith("text/"):
        text = file_bytes.decode("utf-8", errors="replace").strip()
        return ParsedDocument(
            title=filename,
            text=text,
            page_count=1,
            word_count=len(text.split()),
            content_type="text/plain",
        )
    else:
        raise ValueError(f"Unsupported file type: {ext} ({content_type})")
